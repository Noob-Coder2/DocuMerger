import os
import subprocess
import tempfile
import io
import shutil
from typing import List, Union, Tuple
from pypdf import PdfWriter, PdfReader
from docx import Document
from docxcompose.composer import Composer
import tiktoken

# Model configurations with context limits
MODEL_CONFIGS = {
    "GPT-4o": {"limit": 128000, "encoding": "cl100k_base", "icon": "🟢"},
    "GPT-4 Turbo": {"limit": 128000, "encoding": "cl100k_base", "icon": "🔵"},
    "Claude 3.5 Sonnet": {"limit": 200000, "encoding": "cl100k_base", "icon": "🟠"},
    "Gemini 1.5 Pro": {"limit": 1000000, "encoding": "cl100k_base", "icon": "🔴"},
    "Llama 3.3 70B": {"limit": 128000, "encoding": "cl100k_base", "icon": "🟣"},
}

class DocumentProcessor:
    @staticmethod
    def get_token_count(text_content: str, encoding_name: str = "cl100k_base") -> int:
        """Counts tokens using tiktoken."""
        try:
            encoding = tiktoken.get_encoding(encoding_name)
            return len(encoding.encode(text_content))
        except Exception as e:
            print(f"Error counting tokens: {e}")
            return 0
    
    @staticmethod
    def get_model_token_info(text_content: str, model_name: str = "GPT-4o") -> dict:
        """Returns token count and model context info."""
        config = MODEL_CONFIGS.get(model_name, MODEL_CONFIGS["GPT-4o"])
        token_count = DocumentProcessor.get_token_count(text_content, config["encoding"])
        context_limit = config["limit"]
        percentage = (token_count / context_limit) * 100 if context_limit > 0 else 0
        
        return {
            "tokens": token_count,
            "limit": context_limit,
            "percentage": min(percentage, 100),
            "icon": config["icon"],
            "model": model_name
        }

    @staticmethod
    def sanitize_content(text: str, remove_comments: bool = False) -> str:
        """Sanitizes text by removing API keys and optionally comments."""
        import re
        
        # 1. Remove common API Key patterns
        # OpenAI keys
        text = re.sub(r'sk-[a-zA-Z0-9]{20,}', '[REDACTED_API_KEY]', text)
        # AWS Access Keys
        text = re.sub(r'AKIA[0-9A-Z]{16}', '[REDACTED_AWS_KEY]', text)
        # Google API Keys
        text = re.sub(r'AIza[0-9A-Za-z\-_]{35}', '[REDACTED_GOOGLE_KEY]', text)
        # Generic API key patterns (api_key=, apikey=, api-key=)
        text = re.sub(r'(?i)(api[_-]?key\s*[:=]\s*["\']?)([a-zA-Z0-9_\-]{20,})(["\']?)', r'\1[REDACTED_KEY]\3', text)
        
        # 2. Remove comments if requested (Python/JS style)
        if remove_comments:
            # WARNING: Comment removal is heuristic and may not be perfect
            # It tries to preserve URLs and strings but may have edge cases
            lines = text.split('\n')
            cleaned_lines = []
            for line in lines:
                cleaned_line = line
                
                # Skip lines that are mostly URLs (to protect https://)
                if 'http://' in line or 'https://' in line:
                    # Only remove trailing comments after code on these lines
                    # Find # or // that's NOT part of a URL
                    # Simple approach: if there's a space before # or //, it's likely a comment
                    for pattern in [r'\s+#(?!.*://)', r'\s+//(?!.*:)']:
                        match = re.search(pattern, cleaned_line)
                        if match:
                            cleaned_line = cleaned_line[:match.start()]
                else:
                    # For non-URL lines, remove comments normally
                    # Python style comments
                    if '#' in cleaned_line:
                        # Don't remove if it's in a string (basic check for quotes)
                        quote_count = cleaned_line[:cleaned_line.index('#')].count('"') + cleaned_line[:cleaned_line.index('#')].count("'")
                        if quote_count % 2 == 0:  # Even number of quotes = not in string
                            cleaned_line = cleaned_line[:cleaned_line.index('#')]
                    
                    # JS/C++ style comments
                    if '//' in cleaned_line and cleaned_line.count('//') == cleaned_line.count(' //'):
                        # Only remove if // has space before it (likely a comment)
                        parts = cleaned_line.split(' //', 1)
                        if len(parts) > 1:
                            cleaned_line = parts[0]
                
                cleaned_lines.append(cleaned_line.rstrip())
            text = '\n'.join(cleaned_lines)
            
        return text

    @staticmethod
    def merge_as_text(file_list, sanitize: bool = False, remove_comments: bool = False) -> str:
        """
        Text Lane: Instant text merging with Smart Markdown and Sanitization.
        """
        master_string = ""
        for file in file_list:
            filename = file.name
            ext = os.path.splitext(filename)[1].lower()
            
            separator = f"\n\n========================================\n# File: {filename}\n========================================\n\n"
            
            content = ""
            try:
                # Always reset file pointer before reading
                file.seek(0)
                
                # Handle bytes vs string (if file was already read or is a buffer)
                if hasattr(file, 'getvalue'): # BytesIO/StringIO
                    bytes_data = file.read()
                elif hasattr(file, 'read'):
                    bytes_data = file.read()
                else: 
                     bytes_data = b"" # Should not happen

                if isinstance(bytes_data, str):
                    content = bytes_data
                else:
                    try:
                        content = bytes_data.decode('utf-8')
                    except UnicodeDecodeError:
                        content = bytes_data.decode('latin-1', errors='replace')
            except Exception as e:
                content = f"[Error reading file: {str(e)}]"

            # --- Smart Markdown Wrapping ---
            # If it looks like code, wrap it
            code_extensions = {
                '.py': 'python', '.js': 'javascript', '.tsx': 'typescript', 
                '.jsx': 'javascript', '.ts': 'typescript', '.html': 'html', 
                '.css': 'css', '.json': 'json', '.sql': 'sql', '.java': 'java', 
                '.c': 'c', '.cpp': 'cpp', '.sh': 'bash', '.yml': 'yaml', '.yaml': 'yaml'
            }
            
            if ext in code_extensions:
                lang = code_extensions[ext]
                content = f"```{lang}\n{content}\n```"
            
            # --- Sanitization ---
            if sanitize:
                content = DocumentProcessor.sanitize_content(content, remove_comments)

            master_string += separator + content
        
        return master_string

    @staticmethod
    def convert_to_pdf(input_path: str, output_folder: str) -> str:
        """Converts a file to PDF using LibreOffice (Headless)."""
        # Check if soffice is available
        soffice_cmd = "soffice"
        if os.name == 'nt': # Windows
             # Try to find common paths or assume it's in PATH
             # If not in path, this check specifically helps to fail fast
             if shutil.which("soffice") is None:
                  # Try some common windows paths just in case or rely on user PATH
                  pass 

        try:
            # soffice --headless --convert-to pdf --outdir <dir> <file>
            result = subprocess.run(
                [soffice_cmd, '--headless', '--convert-to', 'pdf', '--outdir', output_folder, input_path],
                capture_output=True,
                text=True,
                check=True
            )
            
            # The output filename matches the input basename but with .pdf
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            expected_output = os.path.join(output_folder, base_name + ".pdf")
            
            if os.path.exists(expected_output):
                return expected_output
            else:
                raise FileNotFoundError(f"Conversion failed, PDF not found: {result.stderr}")

        except FileNotFoundError:
             raise RuntimeError("LibreOffice not found. Please ensure 'soffice' is installed and in your PATH.")
        except subprocess.CalledProcessError as e:
            raise RuntimeError(f"LibreOffice conversion error: {e.stderr}")
        except Exception as e:
            raise RuntimeError(f"Unexpected error during conversion: {str(e)}")

    @staticmethod
    def merge_files(file_list, output_format: str, sanitize: bool = False, remove_comments: bool = False) -> Tuple[Union[str, io.BytesIO], str]:
        """
        Routes the files to the appropriate "Lane" based on output_format and input types.
        Returns: (Result Data, MIME Type) or raises Exception
        """
        
        # --- Text Lane ---
        if output_format == "TXT":
            return DocumentProcessor.merge_as_text(file_list, sanitize, remove_comments), "text/plain"

        # --- Word Lane ---
        if output_format == "DOCX":
            # Check if all inputs are docx
            if all(f.name.lower().endswith('.docx') for f in file_list):
                try:
                    merged_doc = Document()
                    composer = Composer(merged_doc)
                    
                    # We need to write streamlit uploads to temp files for docxcompose (it likes file paths or robust streams)
                    # Actually docx.Document() accepts streams.
                    # But docxcompose appends to the first one. Let's create a fresh one.
                    # Use the first file as the master template if possible to keep styles,
                    # but for now let's just create a new empty one and append all? 
                    # docxcompose usually requires a master document.
                    
                    # Strategy: Load first file as master, then append rest.
                    first_file = file_list[0]
                    first_file.seek(0)
                    master_doc = Document(first_file)
                    composer = Composer(master_doc)
                    
                    for i, file in enumerate(file_list):
                        if i == 0: continue # Skip first (master)
                        file.seek(0)
                        sub_doc = Document(file)
                        composer.append(sub_doc)
                    
                    out_stream = io.BytesIO()
                    composer.save(out_stream)
                    out_stream.seek(0)
                    return out_stream, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    
                except Exception as e:
                    raise RuntimeError(f"Error in Word Lane merge: {str(e)}")
            else:
                raise ValueError("DOCX output currently only supports .docx inputs. Please convert other files first or choose PDF output.")

        # --- PDF Output (Fast Lane vs Heavy Lane) ---
        if output_format == "PDF":
            
            # Check for Fast Lane eligibility
            all_pdfs = all(f.name.lower().endswith('.pdf') for f in file_list)
            
            if all_pdfs:
                # --- Fast Lane ---
                try:
                    merger = PdfWriter()
                    for file in file_list:
                        file.seek(0)
                        merger.append(file)
                    
                    out_stream = io.BytesIO()
                    merger.write(out_stream)
                    out_stream.seek(0)
                    return out_stream, "application/pdf"
                except Exception as e:
                   raise RuntimeError(f"Error in Fast Lane (PDF) merge: {str(e)}")
            
            else:
                # --- Heavy Lane ---
                # Requires LibreOffice
                # Workflow:
                # 1. Create temp dir.
                # 2. Save all non-PDF uploads to temp dir.
                # 3. Convert them to PDF.
                # 4. Merge all (original PDFs + converted PDFs).
                
                with tempfile.TemporaryDirectory() as temp_dir:
                    pdf_paths = []
                    
                    for file in file_list:
                        file.seek(0) # Reset pointer
                        ext = os.path.splitext(file.name)[1].lower()
                        
                        if ext == '.pdf':
                            # It's already a PDF, save it directly to temp to handle uniformly by path
                            # OR keep in memory. Mixing paths and streams in pypdf is fine.
                            # Let's save to disk to keep the Heavy Lane uniform and robust.
                            temp_path = os.path.join(temp_dir, file.name)
                            with open(temp_path, "wb") as f:
                                f.write(file.read())
                            pdf_paths.append(temp_path)
                            
                        else:
                            # It's something else (docx, txt, etc.)
                            # Save to temp
                            temp_input_path = os.path.join(temp_dir, file.name)
                            with open(temp_input_path, "wb") as f:
                                f.write(file.read())
                            
                            # Convert
                            try:
                                converted_pdf_path = DocumentProcessor.convert_to_pdf(temp_input_path, temp_dir)
                                pdf_paths.append(converted_pdf_path)
                            except Exception as e:
                                # Start a text file for error log if conversion fails?
                                # For now, fail hard or skip? Plan said "skip it and add note" for image->txt, 
                                # but for PDF heavy lane, if conversion fails, the document is lost.
                                # Let's raise error for now as it's safer than silent loss.
                                raise RuntimeError(f"Failed to convert {file.name}: {str(e)}")

                    # Merge all gathered PDFs
                    try:
                        merger = PdfWriter()
                        for path in pdf_paths:
                            merger.append(path)
                        
                        out_stream = io.BytesIO()
                        merger.write(out_stream)
                        out_stream.seek(0)
                        return out_stream, "application/pdf"
                    except Exception as e:
                        raise RuntimeError(f"Error in Heavy Lane (PDF) final merge: {str(e)}")

        raise ValueError(f"Unsupported output format: {output_format}")
